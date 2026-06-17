"""Convert markdown review document to compressed two-column DOCX."""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_FILE = Path(r"D:\StudyWorks\3.2\机器视觉\课程ppt\docs\机器视觉开卷考试复习资料.md")
OUT_FILE = Path(r"D:\StudyWorks\3.2\机器视觉\课程ppt\docs\机器视觉开卷考试复习资料.docx")

# ── Color map for heading levels ──
HEADING_COLORS = {
    1: "8B0000",  # Dark red
    2: "003366",  # Navy blue
    3: "006666",  # Teal
    4: "4B0082",  # Dark purple
    5: "8B4513",  # Dark brown
}
FONT_NAME = "Microsoft YaHei"
FONT_SIZE = Pt(6)  # 6pt
CELL_PADDING = Pt(1)

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(0.4)
    section.bottom_margin = Cm(0.4)
    section.left_margin = Cm(0.4)
    section.right_margin = Cm(0.4)
    # Two columns with separator
    cols = section._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")} w:equalWidth="1" w:num="2" w:space="108" w:sep="1"/>')
        section._sectPr.append(cols)

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = FONT_SIZE
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.0
# Set East Asian font
rPr = style.element.find(qn('w:rPr'))
if rPr is None:
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    style.element.append(rPr)
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
    rPr.append(rFonts)
else:
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

# Heading styles
for i in range(1, 6):
    hstyle = doc.styles[f'Heading {i}']
    hstyle.font.name = FONT_NAME
    hstyle.font.size = FONT_SIZE
    hstyle.font.bold = True
    hstyle.font.color.rgb = RGBColor.from_string(HEADING_COLORS[i])
    hstyle.paragraph_format.space_before = Pt(0)
    hstyle.paragraph_format.space_after = Pt(0)
    hstyle.paragraph_format.line_spacing = 1.0
    hrPr = hstyle.element.find(qn('w:rPr'))
    if hrPr is None:
        hrPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        hstyle.element.append(hrPr)
    hrFonts = hrPr.find(qn('w:rFonts'))
    if hrFonts is None:
        hrFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        hrPr.append(hrFonts)
    else:
        hrFonts.set(qn('w:eastAsia'), FONT_NAME)

def add_run(para, text, bold=False, italic=False, color=None, size=None):
    """Add a run with formatting."""
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # Set East Asian font on run
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._r.insert(0, rPr)
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        rFonts_elem = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rPr.append(rFonts_elem)
    else:
        rFonts_elem.set(qn('w:eastAsia'), FONT_NAME)
    return run

def add_paragraph(text="", style_name=None, bold=False, color=None):
    """Add paragraph with optional style."""
    if style_name:
        para = doc.add_paragraph(style=style_name)
    else:
        para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    if text:
        add_run(para, text, bold=bold, color=color)
    return para

def parse_inline(text, para):
    """Parse inline markdown: bold, italic, inline math, inline code."""
    # Bold: **text**
    # Italic: *text* (but not **)
    # Inline math: $...$
    # Inline code: `...`
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)|'        # bold
        r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|'  # italic
        r'(\$\$?(.+?)\$\$?)|'     # math display or inline
        r'(`(.+?)`)'               # code
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            add_run(para, text[last:m.start()])
        if m.group(2):  # bold
            add_run(para, m.group(2), bold=True)
        elif m.group(3):  # italic
            add_run(para, m.group(3), italic=True)
        elif m.group(5):  # math
            add_run(para, m.group(5), italic=True, color="336699")
        elif m.group(7):  # code
            add_run(para, m.group(7), color="CC0000")
        last = m.end()
    if last < len(text):
        add_run(para, text[last:])

def process_table(lines_iter, lines_list, idx):
    """Process a markdown table. Returns list of rows, last index."""
    rows = []
    # Collect header + separator + body
    i = idx
    while i < len(lines_list):
        line = lines_list[i]
        if line.startswith('|') and '|' in line[1:]:
            rows.append(line)
            i += 1
        else:
            break

    if len(rows) < 2:
        return i

    # Parse rows
    parsed_rows = []
    for r_idx, row in enumerate(rows):
        cells = [c.strip() for c in row.strip('|').split('|')]
        parsed_rows.append(cells)

    if len(parsed_rows) < 2:
        return i

    header = parsed_rows[0]
    # Skip separator row (contains :--- patterns)
    separator = parsed_rows[1]
    body = parsed_rows[2:] if len(parsed_rows) > 2 else []

    # Check if this is actually a separator row
    is_sep = all(re.match(r'^:?-{3,}:?$', c) for c in separator)
    if not is_sep:
        body = parsed_rows[1:]
        sep_idx = None
    else:
        sep_idx = 1

    # Calculate column count
    ncols = max(len(header), max((len(r) for r in body), default=0))

    # Create table
    table = doc.add_table(rows=1 + len(body), cols=ncols)
    table.style = 'Table Grid'

    # Set table width to full column width
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="auto"/>')
    tblPr.append(tblW)

    # Compact cell margins
    for cell in table._tbl.iter_tcs():
        tcPr = cell.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
            cell.insert(0, tcPr)
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="0" w:type="dxa"/>'
            f'<w:left w:w="18" w:type="dxa"/>'
            f'<w:bottom w:w="0" w:type="dxa"/>'
            f'<w:right w:w="18" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    # Header row
    for c_idx, cell_text in enumerate(header):
        if c_idx < ncols:
            cell = table.rows[0].cells[c_idx]
            cell.paragraphs[0].clear()
            parse_inline(cell_text.strip(), cell.paragraphs[0])
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = FONT_SIZE
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            # Header shading
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                cell._tc.insert(0, tcPr)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8ECF0" w:val="clear"/>')
            tcPr.append(shading)

    # Body rows
    for r_idx, row_cells in enumerate(body):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx < ncols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.paragraphs[0].clear()
                parse_inline(cell_text.strip(), cell.paragraphs[0])
                for run in cell.paragraphs[0].runs:
                    run.font.size = FONT_SIZE
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    return i


def convert_md(md_path):
    """Main conversion function."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.split('\n')

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = add_paragraph()
                add_run(p, code_text, color="333333")
                p.paragraph_format.left_indent = Cm(0.3)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            # Add a thin line
            p = add_paragraph()
            pPr = p._p.find(qn('w:pPr'))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                p._p.insert(0, pPr)
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="2" w:space="1" w:color="999999"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        stripped = line.strip()

        # Heading 1: # text
        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = add_paragraph(stripped[2:], style_name='Heading 1')
            i += 1
            continue

        # Heading 2: ## text
        if stripped.startswith('## ') and not stripped.startswith('### '):
            p = add_paragraph(stripped[3:], style_name='Heading 2')
            i += 1
            continue

        # Heading 3: ### text
        if stripped.startswith('### ') and not stripped.startswith('#### '):
            p = add_paragraph(stripped[4:], style_name='Heading 3')
            i += 1
            continue

        # Heading 4: #### text
        if stripped.startswith('#### '):
            p = add_paragraph(stripped[5:], style_name='Heading 4')
            i += 1
            continue

        # Block quote: > text
        if stripped.startswith('> '):
            p = add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            parse_inline(stripped[2:], p)
            i += 1
            continue

        # Table: | ... |
        if stripped.startswith('|') and '|' in stripped[1:]:
            i = process_table(lines, lines, i)
            continue

        # Bullet list: - text
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.first_line_indent = Cm(-0.15)
            add_run(p, "• ", bold=True)
            parse_inline(stripped[2:], p)
            i += 1
            continue

        # Numbered list: N. text
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            p = add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.first_line_indent = Cm(-0.15)
            add_run(p, f"{num_match.group(1)}. ", bold=True)
            parse_inline(num_match.group(2), p)
            i += 1
            continue

        # Regular paragraph
        p = add_paragraph()
        parse_inline(stripped, p)
        i += 1

    # ── Save ──
    doc.save(str(OUT_FILE))
    print(f"Saved: {OUT_FILE}")

if __name__ == '__main__':
    convert_md(MD_FILE)
