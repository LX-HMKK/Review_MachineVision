import os
import pypandoc
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

# Paths
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
md_file = os.path.join(BASE_DIR, '试卷', '试卷推测补全版.md')
docx_file = os.path.join(BASE_DIR, '试卷', '机器视觉_课程试卷_推测补全版.docx')

def main():
    print(f"Converting {md_file} to {docx_file}...")
    
    # 1. Convert Markdown to Docx using Pandoc
    # This ensures LaTeX math is natively converted to OMML (Word Math)
    pypandoc.convert_file(
        md_file,
        'docx',
        outputfile=docx_file,
        extra_args=['--wrap=none']
    )

    # 2. Post-process the generated Docx
    doc = docx.Document(docx_file)

    # Set up page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Process styles for default font (Normal)
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Post-process paragraphs to highlight specific words and apply fonts
    for p in doc.paragraphs:
        # Reset line spacing to be compact
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        
        # Iterate through runs to colorize '【补】' and '【推测补全】'
        for run in p.runs:
            # Just setting East Asia font for every run to ensure Chinese characters render via YaHei correctly
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            if '【补】' in run.text or '【推测补全】' in run.text:
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                
    # Post-process tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for run in p.runs:
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                        if '【补】' in run.text or '【推测补全】' in run.text:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.bold = True

    doc.save(docx_file)
    print("Conversion completed successfully!")

if __name__ == "__main__":
    main()
